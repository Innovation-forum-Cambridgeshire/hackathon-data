#!/usr/bin/env python3
"""
Generate the Challenge 04 governance Word documents.

    python scripts/build-governance-docs.py --out "<01 Legal & Compliance path>"

Produces:
  1. Notification letter — Monitoring Officer, Huntingdonshire District Council
  2. Notification letter — Huntingdon Town Council
  3. Notification letter — St Ives Town Council
  4. Notification letter — Ramsey Town Council
  5. Method statement — defamation, data protection and the synthetic-corpus decision
  6. Participant output rules — what teams may and may not submit

WHY THE DISTRICT COUNCIL IS ON THIS LIST
    Town and parish councils in England do not hold their own register of members'
    interests. The PRINCIPAL AUTHORITY's Monitoring Officer establishes and maintains it
    on their behalf — for Huntingdon, St Ives and Ramsey that is Huntingdonshire District
    Council, whose Head of Legal Services is the Monitoring Officer. Writing only to the
    town councils would miss the officer who actually holds the register.

CONTACT DETAILS
    Taken from the councils' own published pages (verified 2026-08-15). Named post-holders
    change, so every letter is addressed to the ROLE and carries the name only as a
    courtesy line to be re-checked before sending. Anything unverified is a visible
    [PLACEHOLDER], following the convention already used in the legal pack.
"""

from __future__ import annotations

import argparse
import sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor
except ImportError:  # pragma: no cover
    sys.exit("python-docx is required: pip install python-docx")

TODAY = date.today().strftime("%d %B %Y")

# Registered particulars verified against Companies House 2026-08-15:
# GLOBAL INNOVATION FORUM LIMITED, company no. 10010132, incorporated 17 February 2016,
# private company limited by guarantee, status Active.
# These supersede the [PLACEHOLDER] entries used in earlier drafts.
SENDER = [
    "Global Innovation Forum Limited",
    "Future Business Centre",
    "Kings Hedges Road",
    "Cambridge",
    "CB4 2HY",
    "",
    "info@inno-forum.co.uk",
    "Company number 10010132",
]

# Verified 2026-08-15 from each council's published contact page.
RECIPIENTS = [
    dict(
        key="hdc-monitoring-officer",
        role="The Monitoring Officer (Head of Legal Services)",
        body="Huntingdonshire District Council",
        addr=["Pathfinder House", "St Mary's Street", "Huntingdon", "PE29 3TN"],
        note="Copied to the Data Protection Officer, infogov@3csharedservices.org",
        holds_register=True,
    ),
    dict(
        key="huntingdon-town-council",
        role="The Town Clerk",
        body="Huntingdon Town Council",
        addr=["Town Hall", "Market Hill", "Huntingdon", "PE29 3PJ"],
        note="Current Town Clerk published as Philip Peacock — confirm before sending. "
             "town.council@huntingdontown.gov.uk",
        holds_register=False,
    ),
    dict(
        key="st-ives-town-council",
        role="The Town Clerk",
        body="St Ives Town Council",
        addr=["Town Hall", "Market Hill", "The Old Riverport", "St Ives", "PE27 5AL"],
        note="Town Clerk not named on the published contact page — confirm before sending. "
             "Telephone 01480 388929",
        holds_register=False,
    ),
    dict(
        key="ramsey-town-council",
        role="The Town Clerk",
        body="Ramsey Town Council",
        addr=["Ramsey Abbey Estate Offices", "7 Church Green", "Ramsey", "PE26 1DW"],
        note="Current Town Clerk published as Lisa Renfree — confirm before sending. "
             "ramseytc@ramseytowncouncil.gov.uk",
        holds_register=False,
    ),
]


def base_doc() -> Document:
    d = Document()
    st = d.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)
    return d


def para(d, text, *, bold=False, italic=False, size=None, space_after=6, align=None):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    if size:
        r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    return p


def draft_banner(d, text):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(14)


def notification_letter(rec: dict) -> Document:
    d = base_doc()
    draft_banner(d, "DRAFT — for legal review before sending. Confirm the named post-holder is current.")

    for line in SENDER:
        para(d, line, space_after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    para(d, "", space_after=10)
    para(d, TODAY, space_after=14, align=WD_ALIGN_PARAGRAPH.RIGHT)

    para(d, rec["role"], space_after=0)
    para(d, rec["body"], space_after=0)
    for line in rec["addr"]:
        para(d, line, space_after=0)
    para(d, "", space_after=14)

    para(d, "Dear Sir or Madam", space_after=12)
    para(d,
         "Innovation Forum hackathon challenge “Safe in the Open” — notification of approach, "
         "and confirmation that no member data will be processed",
         bold=True, space_after=12)

    para(d,
         "I am writing to inform you of a forthcoming public hackathon challenge concerning the "
         "online abuse of people in public office, and to confirm the decisions we have taken so "
         "that no personal data relating to your members is processed. We are not seeking access "
         "to any council system, and we are not asking you to provide anything.",
         space_after=10)

    para(d, "What the challenge is", bold=True, space_after=6)
    para(d,
         "Innovation Forum runs a programme of public hackathons in partnership with R1X as its "
         "technology partner. One challenge, “Safe in the Open”, asks teams to prototype tools that "
         "help people in public office who face online abuse — to recognise patterns, signpost "
         "support, and fail safe. The Local Government Association reported in 2025 that 72% of "
         "councillors had experienced abuse.",
         space_after=10)

    para(d, "What data we will use — and what we will not", bold=True, space_after=6)
    para(d,
         "Teams will work on a SYNTHETIC dataset generated by us. It matches the shape of a "
         "published register of members' interests — ward, role, committee, declaration category — "
         "but is not derived from any real person. No real member is described, and nothing "
         "identifying a real member is published.",
         space_after=8)

    if rec["holds_register"]:
        para(d,
             "We are aware that, as the principal authority, your Monitoring Officer establishes and "
             "maintains the register of members' interests on behalf of the town and parish councils "
             "in the district. We want to be explicit that we will not be mirroring, restructuring or "
             "republishing that register. Although it is published under the Open Government Licence, "
             "aggregating public records into a structured, queryable dataset is a distinct processing "
             "operation with its own risk profile, and we have chosen not to undertake it.",
             space_after=8)
    else:
        para(d,
             "We understand that the register of members' interests for your council is maintained by "
             "the Monitoring Officer at Huntingdonshire District Council, whom we are also notifying. "
             "We will not be mirroring, restructuring or republishing that register.",
             space_after=8)

    para(d,
         "We will not collect or republish social media content, local news comments, or any other "
         "material authored by or about identifiable individuals. Where a member has had a sensitive "
         "interest withheld from publication under section 32 of the Localism Act 2011, no attempt "
         "will be made to infer it: that provision exists to protect people at risk of intimidation, "
         "and defeating it would be contrary to the purpose of the challenge.",
         space_after=10)

    para(d, "Why we are writing", bold=True, space_after=6)
    para(d,
         "Because we are using synthetic data only, we do not believe any notification obligation "
         "arises. We are writing as a matter of good practice. You may hold context we could not "
         "reasonably know — a safeguarding matter, or a member with particular concerns — and we "
         "would rather hear it before the event than after.",
         space_after=10)

    para(d,
         "If you have any concern about this challenge, or would prefer we adjust its framing, "
         "please contact us at info@inno-forum.co.uk. If it would be helpful, we are glad to share "
         "the challenge brief and our data protection assessment.",
         space_after=12)

    para(d, "Yours faithfully", space_after=24)
    para(d, "[NAME]", space_after=0)
    para(d, "[ROLE], Innovation Forum", space_after=14)

    if rec.get("note"):
        para(d, f"Note for sender: {rec['note']}", italic=True, size=9)
    return d


def method_statement() -> Document:
    d = base_doc()
    draft_banner(d, "DRAFT — for legal review. Internal method statement; not for publication as-is.")

    para(d, "Challenge 04 “Safe in the Open” — data protection and defamation method statement",
         bold=True, size=14, space_after=12)
    para(d, f"Innovation Forum × R1X · {TODAY}", italic=True, space_after=14)

    para(d, "1. Summary of approach", bold=True, space_after=6)
    para(d,
         "Challenge 04 uses a synthetic corpus. No real person's data is collected, processed or "
         "published. This decision was taken because the alternative — republishing real registers "
         "of interests and real abusive messages — carries legal exposure that a lawful basis under "
         "UK GDPR would not cure, and because the challenge loses nothing by using synthetic data.",
         space_after=10)

    para(d, "2. What the law required us to consider", bold=True, space_after=6)
    for heading, text in [
        ("UK GDPR / Data Protection Act 2018",
         "Councillors are identifiable living individuals. Public availability of a register does not "
         "remove it from the regime; it changes which lawful basis is available. Processing would have "
         "required a legitimate interests assessment, an Article 9 condition for political opinions, "
         "Article 14 transparency to every member, and a DPIA."),
        ("Article 8 ECHR — aggregation",
         "In Catt v United Kingdom (ECtHR, 24 January 2019) the systematic collection and retention of "
         "information gathered entirely from public sources was held to engage Article 8. Turning "
         "scattered public records into a queryable dataset is a new processing operation."),
        ("Defamation Act 2013",
         "Any output implying something about an identifiable office-holder — a “most-targeted member” "
         "ranking, or anything readable as an allegation — could be defamatory. A lawful basis under "
         "UK GDPR is no defence to a defamation claim; they are separate questions."),
        ("Protection from Harassment Act 1997",
         "Republishing abusive messages can form part of a course of conduct, independently of who "
         "originally sent them."),
        ("Misuse of private information",
         "A tort in its own right, engaged by aggregation even where each element was public."),
        ("Communications offences",
         "Online Safety Act 2023 ss.179–181 and Communications Act 2003 s.127. Distributing threatening "
         "or false communications is not automatically safe because someone else sent them first."),
        ("Localism Act 2011 s.32",
         "Members may have a sensitive interest withheld from the published register where disclosure "
         "risks violence or intimidation. Inferring those gaps would defeat a statutory safeguard."),
        ("Platform terms",
         "X, Facebook and Reddit terms generally prohibit redistributing scraped content. This is a "
         "contractual constraint that applies regardless of the legal position."),
    ]:
        para(d, heading, bold=True, space_after=2)
        para(d, text, space_after=8)

    para(d, "3. The Online Safety Act does not apply", bold=True, space_after=6)
    para(d,
         "OSA 2023 duties attach to user-to-user services and search services. The hackathon data "
         "service publishes static files; participants do not upload or share content with one another "
         "through it. It is therefore neither. This conclusion does not depend on whether the corpus is "
         "real or synthetic — but the offences at 2 above do, which is why the corpus is synthetic.",
         space_after=10)

    para(d, "4. Controls implemented", bold=True, space_after=6)
    for text in [
        "Synthetic corpus — no real register, no real messages, generated from a seed held in the repository so results are reproducible for judging.",
        "Real sources are catalogued as reference-only: teams may read published registers at source, but we do not mirror, restructure or republish them.",
        "Only the gold layer is published. Bronze and silver never leave the pipeline.",
        "Anonymisation is irreversible: no surrogate mapping survives a build, so nothing can be re-identified from what we publish.",
        "k-anonymity is enforced as a build gate (k ≥ 5; k ≥ 10 for health-adjacent data), with generalisation preferred over suppression, and evidence written to the Data Asset Register.",
        "A machine-checked catalogue gate blocks any challenge that describes identifiable individuals without either a synthetic alternative or a recorded justification naming a DPIA and approver.",
        "Explicit prohibition on inferring interests withheld under Localism Act 2011 s.32.",
    ]:
        d.add_paragraph(text, style="List Bullet")
    para(d, "", space_after=8)

    para(d, "5. Residual risk", bold=True, space_after=6)
    para(d,
         "Participants may bring their own data and their own models, and we cannot control what they "
         "collect outside the platform. This is addressed by the participant output rules issued with "
         "the challenge brief, by mentor supervision on the day, and by judging criteria that reward "
         "naming the limits of the data rather than overreaching. It cannot be eliminated.",
         space_after=10)

    para(d, "6. Review", bold=True, space_after=6)
    para(d, "Owner: [NAME]. Reviewed by: [LEGAL REVIEWER]. Date: [DATE]. "
            "Next review: before each running of challenge 04.", space_after=6)
    return d


def participant_rules() -> Document:
    d = base_doc()
    draft_banner(d, "DRAFT — to be issued with the challenge brief.")

    para(d, "Challenge 04 “Safe in the Open” — rules for what you may build and submit",
         bold=True, size=14, space_after=12)

    para(d,
         "This challenge is about protecting people. The rules below exist so that the work does not "
         "cause the harm it is trying to reduce. They are judging criteria, not paperwork: entries "
         "that break them cannot be scored.",
         space_after=12)

    para(d, "The data we give you", bold=True, space_after=6)
    para(d,
         "Everything in the challenge dataset is synthetic. The register and the messages were "
         "generated by us. They are realistic in shape but describe nobody. You can therefore analyse, "
         "publish, screenshot and demo any of it freely.",
         space_after=10)

    para(d, "What you must not do", bold=True, space_after=6)
    for text in [
        "Do not collect social media posts, comments or messages about real people — not manually, not with a scraper, not through an API.",
        "Do not name, picture or identify any real councillor, officer or member of the public in your prototype, slides or demo.",
        "Do not attempt to work out an interest a member has had withheld from a public register. That protection exists because someone may be at risk.",
        "Do not build anything that ranks, scores or profiles real named individuals.",
        "Do not upload the challenge data, or anything derived from it about real people, into a third-party service you do not control.",
    ]:
        d.add_paragraph(text, style="List Bullet")
    para(d, "", space_after=8)

    para(d, "What good looks like", bold=True, space_after=6)
    for text in [
        "Work at the level of patterns and categories, not individuals.",
        "Say what your data cannot tell you. Recorded abuse shows what was reported, not what happened.",
        "Make it fail safe: if your tool is unsure, it should say so rather than guess.",
        "Always leave a route to human help. A tool that spots harm and offers nowhere to turn has not finished the job.",
        "If you use an LLM, assume anything you paste into it leaves your control. Do not paste anything about a real person.",
    ]:
        d.add_paragraph(text, style="List Bullet")
    para(d, "", space_after=8)

    para(d, "If you are unsure", bold=True, space_after=6)
    para(d,
         "Ask a mentor before you build it, not after. “Could we scrape…” is always worth asking out "
         "loud — the answer is usually no, and there is usually a better route to the same insight.",
         space_after=6)
    return d


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--out", required=True, help="output directory")
    args = ap.parse_args()

    out = Path(args.out)
    out.mkdir(parents=True, exist_ok=True)
    written = []

    for rec in RECIPIENTS:
        p = out / f"C04 Notification - {rec['body']}.docx"
        notification_letter(rec).save(p)
        written.append(p)

    p = out / "C04 Method Statement - Data Protection and Defamation.docx"
    method_statement().save(p)
    written.append(p)

    p = out / "C04 Participant Output Rules.docx"
    participant_rules().save(p)
    written.append(p)

    print(f"Wrote {len(written)} document(s) to {out}\n")
    for p in written:
        print(f"  {p.name}  ({p.stat().st_size:,} bytes)")
    print("\n  All marked DRAFT. Confirm named post-holders before sending — clerks change.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
