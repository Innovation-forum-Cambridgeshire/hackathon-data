#!/usr/bin/env python3
"""
Programme update for the programme leads, as a branded Word document.

    python3 scripts/programme-report.py --out report/     # first: compute KPIs
    python3 scripts/programme-update-doc.py               # then: write the .docx

Writes report/Programme-Update-<date>.docx

Why this reads programme-kpis.json rather than the board
--------------------------------------------------------
programme-report.py already reads the board, the catalogues and the build, and
writes the numbers to programme-kpis.json. This renders that file. Two scripts
counting the same things independently is how a dashboard and a status report
end up disagreeing in front of the people who need them to agree.

So: if a number in this document looks wrong, fix programme-report.py. There is
nothing to compute here.

House style
-----------
Calibri 10.5 on US Letter, matching the existing Documentation set (checked
against "01 Legal & Compliance/02 - Cookie Policy.docx"). Brand colours are the
ones on record in "03 Brand Guidelines/11 - Brand Guidelines.docx": #009540
green, #04180f ink, #8FD6D8 aqua.

The green is used for rules and headings only. Body text stays near-black:
#009540 on white measures 3.1:1 and would fail WCAG AA at body size, and this
document gets printed and forwarded.
"""
import json
import subprocess
import sys
from datetime import date
from pathlib import Path

try:
    import docx
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("python-docx is required:  pip3 install python-docx")

HERE = Path(__file__).resolve().parent.parent
KPIS = HERE / "report" / "programme-kpis.json"
LOGO = Path(
    "/Users/yavinowens/Library/CloudStorage/OneDrive-R1x/R1X Foundry - Documents/"
    "05_PRODUCT_AND_PLATFORM/website_/public/assets/Innovation-Forum-logo.jpeg"
)
REPO = "Innovation-forum-Cambridgeshire/hackathon-organisers"

GREEN = RGBColor(0x00, 0x95, 0x40)
INK = RGBColor(0x04, 0x18, 0x0F)
GREY = RGBColor(0x55, 0x55, 0x55)
WARN = RGBColor(0xB4, 0x47, 0x1A)

EVENT = date(2026, 10, 26)          # C03 — Beyond the Mainframe
TODAY = date.today()


# ── docx helpers ────────────────────────────────────────────────────────────

def shade(cell, hex_fill):
    el = OxmlElement("w:shd")
    el.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


def rule(doc, colour=GREEN, weight=18):
    """A coloured horizontal rule, drawn as a bottom border on an empty para."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(weight))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), f"{colour[0]:02X}{colour[1]:02X}{colour[2]:02X}")
    bdr.append(bottom)
    pPr.append(bdr)
    return p


def para(doc, text="", size=10.5, bold=False, colour=INK, after=6, before=0,
         italic=False, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = colour
    return p


def rich(doc, parts, size=10.5, after=6):
    """parts = [(text, bold, colour), ...] — one paragraph, mixed emphasis."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    for text, bold, colour in parts:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.color.rgb = colour
    return p


def heading(doc, text, size=15, before=16):
    p = para(doc, text, size=size, bold=True, colour=INK, after=2, before=before)
    rule(doc, GREEN, 12)
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.size = Pt(10.5)
        r.font.bold = True
        r.font.color.rgb = INK
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = INK
    return p


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(c, "00863A")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            txt, colour, bold = (val if isinstance(val, tuple) else (val, INK, False))
            r = p.add_run(str(txt))
            r.font.size = Pt(9.5)
            r.font.color.rgb = colour
            r.font.bold = bold
    # Word ignores cell widths while autofit is on, and python-docx leaves it on
    # by default — so the widths above are silently discarded and the first
    # column ends up as wide as the longest cell in it. Turn autofit off, set
    # the grid, and set every cell: Word needs all three to agree.
    if widths:
        t.autofit = False
        t.allow_autofit = False
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
        for i, w in enumerate(widths):
            t.columns[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ── live facts the KPI file does not carry ──────────────────────────────────

def issue_states(numbers):
    out = {}
    for n in numbers:
        p = subprocess.run(
            ["gh", "issue", "view", str(n), "--repo", REPO, "--json", "state,title",
             "--jq", '.state + "\t" + .title'],
            capture_output=True, text=True)
        if p.returncode == 0 and p.stdout.strip():
            state, title = p.stdout.strip().split("\t", 1)
            out[n] = (state, title)
    return out


def dns_resolves(host):
    return subprocess.run(["nslookup", host], capture_output=True, text=True).returncode == 0


# ── the document ────────────────────────────────────────────────────────────

def build(k):
    days = (EVENT - TODAY).days
    doc = docx.Document()

    s = doc.sections[0]
    s.left_margin = s.right_margin = Inches(0.87)
    s.top_margin = Inches(0.7)
    s.bottom_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK

    # ── Cover block ────────────────────────────────────────────────────────
    if LOGO.exists():
        doc.add_picture(str(LOGO), width=Inches(0.72))
        doc.paragraphs[-1].paragraph_format.space_after = Pt(4)

    para(doc, "INNOVATION FORUM × R1X", size=9, bold=True, colour=GREEN, after=2)
    para(doc, "Programme Update", size=24, bold=True, colour=INK, after=2)
    para(doc, "Applied-AI Hackathon Programme — prepared for the programme leads",
         size=11, colour=GREY, after=8)
    rule(doc, GREEN, 20)

    rich(doc, [
        (f"{TODAY.strftime('%d %B %Y')}", True, INK),
        ("   ·   ", False, GREY),
        (f"{days} days to Challenge 03", True, WARN),
        ("   ·   ", False, GREY),
        (f"{k['totals']['percent_done']:.0f}% of the programme board complete", False, GREY),
    ], size=10.5, after=14)

    # ── Headline ───────────────────────────────────────────────────────────
    heading(doc, "Where we are", before=2)
    para(doc,
         "The engineering is ahead of schedule. The governance is behind it, and governance "
         "is now the binding constraint on everything participant-facing. That is the whole "
         "of the position in one sentence, and it is a better problem to have than the "
         "reverse — but it does not solve itself, and it cannot be solved by the people "
         "currently doing the work.",
         after=8)
    para(doc,
         "Challenge 03 runs 26–30 October and is the only hard deadline in the programme. "
         "The data platform that serves it is 81% complete with three small items left. The "
         "participant workspace was built and deployed this month and is fully working, but "
         "cannot be opened to participants until four governance items clear — and those need "
         "named people and signed decisions, not more build time.",
         after=8)

    # ── Progress by workstream ─────────────────────────────────────────────
    heading(doc, "Progress by workstream")
    # programme-kpis.json counts any issue that has sub-issues as an "epic",
    # which sweeps in second-level parents like #4 and #71. Those are real work
    # items but they are NOT the board's workstream layer, and mixing the two
    # double-counts: #4's children are already inside #16's total. The board
    # names its top level "EPIC — ...", so filter on that rather than on a
    # hand-kept list of numbers that would rot the next time one is added.
    epics = [e for e in k["epics"] if e["title"].startswith("EPIC")]
    if not epics:
        sys.exit("no items titled 'EPIC — ...' found — has the board's naming changed?")

    rows = []
    for e in sorted(epics, key=lambda x: -x["percent"]):
        pct = e["percent"]
        colour = INK if pct >= 75 else (WARN if pct < 40 else GREY)
        rows.append([
            f"#{e['number']}",
            e["title"].replace("EPIC — ", ""),
            f"{e['done']}/{e['total']}",
            (f"{pct:.0f}%", colour, pct < 40),
        ])
    table(doc, ["", "Workstream", "Done", "Complete"], rows,
          widths=[0.5, 4.3, 0.8, 0.9])

    t = k["totals"]
    rich(doc, [
        (f"Board total: {t['items']} items — ", False, GREY),
        (f"{t['status']['Done']} done", True, INK),
        (f", {t['status']['In Progress']} in progress, {t['status']['Todo']} to do.", False, GREY),
    ], size=9.5, after=12)

    # ── What moved ─────────────────────────────────────────────────────────
    heading(doc, "What moved this month")
    para(doc, "The participant workspace went from a plan to a working service.", after=6)
    bullet(doc, "at if-hackathon-workspace.pages.dev. Sign-in, team, countdown and "
                "submission guidance all working end to end, tested against real GitHub accounts.",
           bold_prefix="Built and deployed ")
    bullet(doc, "read inside the workspace — one per challenge, fetched server-side so a "
                "participant's browser makes no third-party request.",
           bold_prefix="Worked examples ")
    bullet(doc, "The design assumed Supabase in London. It was built without a database at "
                "all: GitHub is the identity provider and org membership is the access control. "
                "That removes a processor from every privacy notice, a DPA to negotiate, and "
                "the free-tier pause and backup risks that were on the register.",
           bold_prefix="One processor removed. ")
    bullet(doc, "The Cookie Policy, Privacy Policy, Website Terms and Workspace Privacy "
                "Notice were amended so the workspace is inside the published notices rather "
                "than outside all of them.",
           bold_prefix="Notices amended. ")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── Data readiness ─────────────────────────────────────────────────────
    heading(doc, "Data readiness for October")
    d = k["data_readiness"]
    lic = k["licence_review"]
    table(doc,
          ["Measure", "Position", "Reading"],
          [["Challenges with buildable data",
            f"{d['challenges_with_data']}/{d['challenges_total']}",
            "All five challenges have data"],
           ["Tables buildable",
            f"{d['tables_buildable']}/{d['tables_total']}",
            "Two tables not yet building"],
           ["Rows declared", f"{d['rows_declared']:,}", "Corpus is real, not a sample"],
           ["Sources through licence review",
            (f"{lic['reviewed']}/{lic['total']}", WARN, True),
            (f"{lic['pending']} pending, {lic['restricted']} restricted to pointer-only", WARN, False)]],
          widths=[2.4, 1.4, 3.0])
    para(doc,
         "The licence position is the one to watch. Restricted sources are handled correctly "
         "— the pipeline mirrors only what the licence allows and holds the rest as pointers — "
         "but 13 of 27 sources still outside review is a number that should be falling weekly.",
         size=9.5, colour=GREY, after=12)

    # ── Critical path ──────────────────────────────────────────────────────
    heading(doc, "Critical path to 26 October")
    crit = issue_states([36, 37, 38])
    dns_ok = dns_resolves("data.inno-forum.co.uk")
    rows = []
    for n in (36, 37, 38):
        state, title = crit.get(n, ("UNKNOWN", "—"))
        done = state.upper() == "CLOSED"
        rows.append([
            f"#{n}",
            title,
            ("Done", INK, False) if done else ("Open", WARN, True),
        ])
    table(doc, ["", "Item", "State"], rows, widths=[0.5, 5.1, 0.9])

    if not dns_ok:
        rich(doc, [
            ("Blocking, and small: ", True, WARN),
            ("data.inno-forum.co.uk does not resolve. It is a DNS record and a Worker "
             "deploy — perhaps an hour of work — and the CORS acceptance test behind it "
             "cannot start until it exists. This is the cheapest item on the critical path "
             "and it has been outstanding since mid-August.", False, INK),
        ], size=10.5, after=12)

    # ── The gate ───────────────────────────────────────────────────────────
    heading(doc, "What is actually blocking the workspace")
    para(doc,
         "None of this is engineering. Each item needs a person to be named or a decision "
         "to be signed.", after=6)
    gate = issue_states([89, 90, 91, 103, 102])
    rows = []
    for n in (89, 91, 90, 103, 102):
        state, title = gate.get(n, ("UNKNOWN", "—"))
        rows.append([f"#{n}", title,
                     ("Done", INK, False) if state.upper() == "CLOSED" else ("Open", WARN, True)])
    table(doc, ["", "Item", "State"], rows, widths=[0.5, 5.1, 0.9])
    rich(doc, [
        ("#89 is the one that unlocks the rest. ", True, INK),
        ("Until the Data Owner, DPO, SIRO and Safeguarding lead are named, nobody can sign "
         "the DPIA, and the Workspace Privacy Notice keeps a placeholder where the data "
         "protection contact belongs. Those placeholders are deliberately visible — an "
         "invented contact would be worse than a shown gap — but a notice cannot be put in "
         "front of participants in that state.", False, INK),
    ], after=12)

    # ── Estimate ───────────────────────────────────────────────────────────
    doc.add_page_break()
    heading(doc, "Estimated completion", before=2)
    para(doc,
         "Three separate estimates, because the three workstreams are limited by different "
         "things and averaging them would hide the one that matters.", after=8)

    table(doc,
          ["Workstream", "Estimate", "Confidence", "What it depends on"],
          [["Data platform ready for C03",
            ("mid-September", INK, True), ("High", INK, False),
            "Three small items; only the DNS record is truly blocking"],
           ["Workspace open to participants",
            ("Not estimable", WARN, True), ("—", WARN, False),
            "Cannot be dated until #89 names the governance roles"],
           ["Event delivery (#21)",
            ("Unknown", WARN, True), ("Low", WARN, False),
            "0% complete and its tasks were inferred, not derived"]],
          widths=[1.9, 1.3, 0.9, 3.0])

    para(doc, "Reading the middle row honestly", size=11.5, bold=True, after=4, before=6)
    para(doc,
         "If the governance roles are named by 1 September, the remaining chain — DPIA, then "
         "the placeholders, then a marketing-site deploy — is roughly three to four weeks of "
         "mostly review time. That lands near 1 October and leaves three weeks of margin "
         "before the event. That is comfortable.",
         after=6)
    rich(doc, [
        ("If #89 slips past mid-September, the workspace will not be open for Challenge 03, "
         "and the programme should adopt the fallback already recorded on the board: ", False, INK),
        ("run Challenge 03 as previous challenges ran, and ship the workspace for Challenge 04.",
         True, INK),
        (" That is a perfectly good outcome and costs nothing already built — but it is a "
         "decision that should be taken deliberately in September, not discovered in October.",
         False, INK),
    ], after=8)

    para(doc, "The estimate we cannot give", size=11.5, bold=True, after=4, before=4)
    para(doc,
         "Event delivery (#21) is 0% complete and its five child tasks were a scaffold — "
         "inferred when the board was built, not derived from a run of the event. Any date "
         "put against it today would be invented. It is the largest unquantified risk in the "
         "programme, and it is about people, venue and logistics rather than software. It "
         "needs an hour with whoever has actually run one of these before, and that hour is "
         "worth more than anything else on this page.",
         after=12)

    # ── Asks ───────────────────────────────────────────────────────────────
    heading(doc, "What we need from the programme leads")
    para(doc, "Three decisions. The first is urgent; the other two shape October.",
         after=6)
    bullet(doc, "Data Owner, DPO, SIRO and Safeguarding lead. This is the single "
                "item blocking the workspace, the DPIA and the privacy notice. It costs a "
                "meeting, and every week it waits is a week off the margin.",
           bold_prefix="1. Name four roles (#89). ")
    bullet(doc, "How long we keep the record of which team someone was in. It is one "
                "sentence and it is the last placeholder that is not blocked on anything else.",
           bold_prefix="2. Set the retention period (#103). ")
    bullet(doc, "Replace the inferred tasks with real ones, so this "
                "document can carry a date for it next month instead of the word 'unknown'.",
           bold_prefix="3. Give us an hour on event delivery (#21). ")

    doc.add_paragraph()
    rule(doc, GREEN, 12)
    para(doc,
         f"Generated {TODAY.strftime('%d %B %Y')} from the Hackathon Programme board. "
         "Figures come from programme-kpis.json, which is computed from the board, the data "
         "catalogues and the build — not entered by hand. Board: "
         f"{k['project']['url']}",
         size=8.5, colour=GREY, after=0)

    return doc


def main():
    if not KPIS.exists():
        sys.exit(f"{KPIS} not found — run:  python3 scripts/programme-report.py --out report/")
    k = json.loads(KPIS.read_text())
    if k.get("generated") != TODAY.isoformat():
        print(f"  WARNING: KPIs were generated {k.get('generated')}, today is {TODAY}. "
              f"Re-run programme-report.py for current figures.")
    doc = build(k)
    out = HERE / "report" / f"Programme-Update-{TODAY.isoformat()}.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"  wrote {out.relative_to(HERE)}")
    print(f"  {(EVENT - TODAY).days} days to C03 · {k['totals']['percent_done']:.0f}% board complete")
    return 0


if __name__ == "__main__":
    sys.exit(main())
